"""Protocols feature - protocol library with manual entry, Claude round-trip import,
recipe tables, and run history."""
from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from pydantic import BaseModel
from typing import Optional, List
from datetime import datetime, timedelta
import json, re, io

from core.database import register_table, register_seed, ensure_column, get_db
from core.llm import fetch_url_text  # HTTP scrape helper, not an LLM call

register_table("protocols", """CREATE TABLE IF NOT EXISTS protocols (
    id          INTEGER PRIMARY KEY AUTOINCREMENT,
    title       TEXT NOT NULL,
    source_type TEXT NOT NULL DEFAULT 'url',
    url         TEXT DEFAULT NULL,
    source_text TEXT DEFAULT NULL,
    steps       TEXT DEFAULT NULL,
    recipe      TEXT DEFAULT NULL,
    notes       TEXT NOT NULL DEFAULT '',
    tags        TEXT NOT NULL DEFAULT '[]',
    created     TEXT NOT NULL,
    updated     TEXT NOT NULL)""")

register_table("active_runs", """CREATE TABLE IF NOT EXISTS active_runs (
    run_id       TEXT PRIMARY KEY,
    protocol_id  INTEGER NOT NULL,
    protocol_json TEXT NOT NULL,
    steps_json   TEXT NOT NULL DEFAULT '[]',
    recipe_json  TEXT DEFAULT NULL,
    group_name   TEXT NOT NULL DEFAULT '',
    subgroup     TEXT NOT NULL DEFAULT '',
    scaling      INTEGER NOT NULL DEFAULT 0,
    scale_factor REAL NOT NULL DEFAULT 1.0,
    started_at   TEXT NOT NULL,
    updated_at   TEXT NOT NULL)""")

# Idempotent column-add for the "Are you still running this?" daily check-in.
# Stores ISO-8601 timestamp until which we should NOT prompt; null/empty = ask.
register_seed(lambda conn: ensure_column(conn, "active_runs",
                                         "snoozed_until", "TEXT DEFAULT NULL"))

register_table("protocol_runs", """CREATE TABLE IF NOT EXISTS protocol_runs (
    id           INTEGER PRIMARY KEY AUTOINCREMENT,
    protocol_id  INTEGER NOT NULL,
    date         TEXT NOT NULL,
    group_name   TEXT NOT NULL DEFAULT '',
    steps_done   INTEGER NOT NULL DEFAULT 0,
    steps_total  INTEGER NOT NULL DEFAULT 0,
    deviations   INTEGER NOT NULL DEFAULT 0,
    steps_json   TEXT DEFAULT NULL,
    recipe_json  TEXT DEFAULT NULL,
    entry_id     INTEGER DEFAULT NULL,
    created      TEXT NOT NULL)""")

def _migrate():
    with get_db() as conn:
        for stmt in [
            "ALTER TABLE protocols ADD COLUMN source_type TEXT NOT NULL DEFAULT 'url'",
            "ALTER TABLE protocols ADD COLUMN recipe TEXT DEFAULT NULL",
            "ALTER TABLE protocols ADD COLUMN auto_complete TEXT NOT NULL DEFAULT 'manual'",
        ]:
            try:
                conn.execute(stmt)
                conn.commit()
            except Exception:
                pass
_migrate()

# --------------------------------------------------------------------------- #
#  Claude round-trip helpers
#
#  Flow: user hits "Copy Claude prompt" on a protocol → clipboard gets the
#  template + source_text below. User pastes into Claude chat, Claude returns
#  a delimited block. User pastes that back via "Import from Claude", which
#  parses steps + recipe tables + notes and updates the protocol.
# --------------------------------------------------------------------------- #

CLAUDE_PROMPT_TEMPLATE = """You are formatting a lab protocol for import into a lab notebook system.
Read the source protocol below and output ONLY the formatted block, using
these exact delimiters:

=== STEPS ===
Numbered list, one step per line. Keep steps atomic and imperative.

=== RECIPES ===
For each table in the protocol (master mix, thermocycler program, buffer,
gel loading, etc.), output a level-2 markdown header with the table name
followed by a standard markdown table. Example:

## Master Mix
| Component | Stock conc. | Volume (uL) | Final conc. |
|-----------|-------------|-------------|-------------|
| Buffer    | 10x         | 2           | 1x          |

=== NOTES ===
Any warnings, tips, or context worth preserving. Free text.

Rules:
- Do not invent volumes, concentrations, temperatures, or times not present
  in the source. If a value is missing, leave the cell blank or write "?".
- Do not add commentary before or after the delimited block.
- If a section has no content, still include the delimiter with nothing under it.

--- SOURCE PROTOCOL ---
"""

def _clean_source_text(text: str, truncate: int = 0) -> str:
    """Strip HTML/entities. Pass truncate>0 to limit length."""
    if not text:
        return ""
    text = re.sub(r'<style[^>]*>.*?</style>', ' ', text, flags=re.DOTALL)
    text = re.sub(r'<script[^>]*>.*?</script>', ' ', text, flags=re.DOTALL)
    text = re.sub(r'<[^>]+>', ' ', text)
    text = re.sub(r'&nbsp;', ' ', text)
    text = re.sub(r'&[a-z]+;', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text[:truncate] if truncate else text


def _split_claude_sections(formatted: str) -> dict:
    """Split a Claude-formatted import block on === DELIMITER === markers.
    Returns dict with keys 'STEPS', 'RECIPES', 'NOTES' (missing sections → '')."""
    # Match "=== NAME ===" at the start of a line
    parts = re.split(r'(?m)^===\s*([A-Z]+)\s*===\s*$', formatted)
    # re.split with a capturing group produces: [pre, name1, body1, name2, body2, ...]
    out = {"STEPS": "", "RECIPES": "", "NOTES": ""}
    for i in range(1, len(parts) - 1, 2):
        name = parts[i].strip().upper()
        body = parts[i + 1].strip()
        if name in out:
            out[name] = body
    return out


def _parse_steps_block(block: str) -> str:
    """Turn a numbered/bulleted list of steps into the JSON shape stored in
    the `steps` column: [{"text": "..."}]."""
    steps = []
    for raw in block.splitlines():
        line = raw.strip()
        # Strip leading "1.", "1)", "- ", "* ", etc.
        line = re.sub(r'^(?:\d+[\.\)]|[-*•])\s*', '', line)
        if line and len(line) > 2:
            steps.append({"text": line})
    return json.dumps(steps)


def _parse_markdown_table(lines: list) -> Optional[dict]:
    """Parse a markdown pipe-table into {columns, rows}. Returns None if the
    lines don't look like a valid table."""
    # Keep only lines that contain a pipe — Claude sometimes leaves blank
    # lines between the table and surrounding text
    rows = [l.strip() for l in lines if '|' in l]
    if len(rows) < 2:
        return None

    def split_row(r):
        # Strip leading/trailing pipes then split
        r = r.strip()
        if r.startswith('|'):
            r = r[1:]
        if r.endswith('|'):
            r = r[:-1]
        return [c.strip() for c in r.split('|')]

    header = split_row(rows[0])
    # rows[1] is usually the |---|---| separator; skip any row that's just
    # dashes/colons/pipes/whitespace
    body_rows = []
    for r in rows[1:]:
        if re.fullmatch(r'[\s\-:|]+', r):
            continue
        cells = split_row(r)
        # Pad short rows, trim long ones, so every row matches header width
        if len(cells) < len(header):
            cells += [''] * (len(header) - len(cells))
        elif len(cells) > len(header):
            cells = cells[:len(header)]
        body_rows.append(cells)

    if not header:
        return None
    return {"columns": header, "rows": body_rows}


def _parse_recipes_block(block: str) -> Optional[str]:
    """Turn a RECIPES block (multiple '## Name' sections each with a markdown
    table) into the JSON stored in the `recipe` column.

    Returns single-table shape {name, columns, rows} if only one table,
    array [{name, columns, rows}, ...] if multiple, or None if empty."""
    if not block.strip():
        return None

    # Split on level-2 headers. Anything before the first '## ' is discarded.
    chunks = re.split(r'(?m)^##\s+(.+?)\s*$', block)
    # chunks = [preamble, name1, body1, name2, body2, ...]
    tables = []
    for i in range(1, len(chunks) - 1, 2):
        name = chunks[i].strip()
        body_lines = chunks[i + 1].splitlines()
        parsed = _parse_markdown_table(body_lines)
        if parsed and parsed["columns"]:
            parsed["name"] = name or "Recipe"
            tables.append(parsed)

    # Fallback: no '## Name' headers, but there might still be a bare markdown
    # table (Claude occasionally omits the header for single-table protocols)
    if not tables:
        parsed = _parse_markdown_table(block.splitlines())
        if parsed and parsed["columns"]:
            parsed["name"] = "Recipe"
            tables.append(parsed)

    if not tables:
        return None
    if len(tables) == 1:
        return json.dumps(tables[0])
    return json.dumps(tables)

async def _text_from_upload(file: UploadFile) -> str:
    content = await file.read()
    fname = (file.filename or "").lower()
    if fname.endswith(".pdf"):
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            return " ".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            raise HTTPException(400, "pypdf not installed")
    if fname.endswith(".docx"):
        try:
            import docx
            from docx.oxml.ns import qn
            doc = docx.Document(io.BytesIO(content))
            lines = []
            for p in doc.paragraphs:
                t = p.text.strip()
                if t:
                    lines.append(t)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            for el in doc.element.body.iter():
                if el.tag == qn('w:txbxContent'):
                    for child in el.iter(qn('w:t')):
                        t = (child.text or "").strip()
                        if t:
                            lines.append(t)
            extracted = "\n".join(lines).strip()
            if not extracted:
                raise HTTPException(400, "Could not extract text from this .docx")
            return extracted
        except HTTPException:
            raise
        except ImportError:
            raise HTTPException(400, "python-docx not installed")
        except Exception as e:
            raise HTTPException(400, "Failed to read .docx: " + str(e))
    return content.decode("utf-8", errors="ignore")

DEFAULT_RECIPE = {
    "columns": ["Component", "Stock conc.", "Volume (uL)", "Final conc."],
    "rows": []
}

# --------------------------------------------------------------------------- #
#  Models
# --------------------------------------------------------------------------- #
class CreateProtocol(BaseModel):
    title: str
    url:   Optional[str] = None
    notes: str = ""
    tags:  List[str] = []
    auto_complete: str = "manual"

class PasteProtocol(BaseModel):
    title: str
    text:  str
    notes: str = ""
    tags:  List[str] = []
    auto_complete: str = "manual"

class ManualProtocol(BaseModel):
    title:  str
    steps:  List[str] = []
    recipe: Optional[str] = None
    notes:  str = ""
    tags:   List[str] = []
    auto_complete: str = "manual"

class UpdateProtocol(BaseModel):
    title:  Optional[str] = None
    notes:  Optional[str] = None
    tags:   Optional[List[str]] = None
    steps:  Optional[str] = None
    recipe: Optional[str] = None
    auto_complete: Optional[str] = None

class ActiveRunCreate(BaseModel):
    run_id:       str
    protocol_id:  int
    protocol_json: str
    steps_json:   str = '[]'
    recipe_json:  Optional[str] = None
    group_name:   str = ''
    subgroup:     str = ''
    scaling:      bool = False
    scale_factor: float = 1.0
    started_at:   str

class ActiveRunUpdate(BaseModel):
    steps_json:   Optional[str] = None
    recipe_json:  Optional[str] = None
    scaling:      Optional[bool] = None
    scale_factor: Optional[float] = None

class SaveRun(BaseModel):
    protocol_id: int
    date:        str
    group_name:  str
    steps_json:  str
    recipe_json: Optional[str] = None
    entry_id:    Optional[int] = None

# --------------------------------------------------------------------------- #
#  Router
# --------------------------------------------------------------------------- #
router = APIRouter(prefix="/api", tags=["protocols"])

@router.get("/protocols")
def list_protocols(tag: str = ""):
    with get_db() as conn:
        rows = conn.execute("SELECT * FROM protocols ORDER BY created DESC").fetchall()
    protocols = [dict(r) for r in rows]
    if tag:
        protocols = [p for p in protocols if tag in json.loads(p.get("tags") or "[]")]
    return {"protocols": protocols}

@router.get("/protocols/{protocol_id}")
def get_protocol(protocol_id: int):
    with get_db() as conn:
        row = conn.execute("SELECT * FROM protocols WHERE id=?", (protocol_id,)).fetchone()
    if not row:
        raise HTTPException(404, "Not found")
    return dict(row)

@router.get("/protocols/{protocol_id}/runs")
def get_protocol_runs(protocol_id: int):
    with get_db() as conn:
        rows = conn.execute(
            "SELECT * FROM protocol_runs WHERE protocol_id=? ORDER BY created DESC",
            (protocol_id,)).fetchall()
    return {"runs": [dict(r) for r in rows]}

@router.post("/protocol-runs")
def save_protocol_run(body: SaveRun):
    now = datetime.utcnow().isoformat()
    steps = json.loads(body.steps_json) if body.steps_json else []
    done  = sum(1 for s in steps if s.get("done"))
    devs  = sum(1 for s in steps if s.get("deviation", "").strip())
    with get_db() as conn:
        cur = conn.execute(
            "INSERT INTO protocol_runs (protocol_id,date,group_name,steps_done,steps_total,"
            "deviations,steps_json,recipe_json,entry_id,created) VALUES (?,?,?,?,?,?,?,?,?,?)",
            (body.protocol_id, body.date, body.group_name, done, len(steps),
             devs, body.steps_json, body.recipe_json, body.entry_id, now))
        conn.commit()
        return dict(conn.execute("SELECT * FROM protocol_runs WHERE id=?", (cur.lastrowid,)).fetchone())


@router.delete("/protocol-runs/{run_id}")
def delete_protocol_run(run_id: int):
    """Remove a protocol_run row. Does NOT delete the linked notebook entry —
    that's the user's data and they can delete it separately if they want.
    Useful for cleaning up accidentally-duplicated runs."""
    with get_db() as conn:
        row = conn.execute("SELECT 1 FROM protocol_runs WHERE id=?", (run_id,)).fetchone()
        if not row:
            raise HTTPException(404, "Protocol run not found")
        conn.execute("DELETE FROM protocol_runs WHERE id=?", (run_id,))
        conn.commit()
    return {"deleted": run_id}



@router.get("/active-runs")
def list_active_runs():
    with get_db() as conn:
        rows = conn.execute(
            "SELECT * FROM active_runs ORDER BY updated_at DESC").fetchall()
    return {"runs": [dict(r) for r in rows]}

@router.post("/active-runs")
def create_active_run(body: ActiveRunCreate):
    now = datetime.utcnow().isoformat()
    with get_db() as conn:
        conn.execute("""INSERT OR REPLACE INTO active_runs
            (run_id,protocol_id,protocol_json,steps_json,recipe_json,
             group_name,subgroup,scaling,scale_factor,started_at,updated_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?)""",
            (body.run_id, body.protocol_id, body.protocol_json, body.steps_json,
             body.recipe_json, body.group_name, body.subgroup,
             1 if body.scaling else 0, body.scale_factor, body.started_at, now))
        conn.commit()
    return {"run_id": body.run_id}

@router.put("/active-runs/{run_id}")
def update_active_run(run_id: str, body: ActiveRunUpdate):
    now = datetime.utcnow().isoformat()
    with get_db() as conn:
        row = conn.execute(
            "SELECT * FROM active_runs WHERE run_id=?", (run_id,)).fetchone()
        if not row:
            raise HTTPException(404, "Active run not found")
        r = dict(row)
        if body.steps_json   is not None: r["steps_json"]   = body.steps_json
        if body.recipe_json  is not None: r["recipe_json"]  = body.recipe_json
        if body.scaling      is not None: r["scaling"]      = 1 if body.scaling else 0
        if body.scale_factor is not None: r["scale_factor"] = body.scale_factor
        conn.execute("""UPDATE active_runs SET
            steps_json=?, recipe_json=?, scaling=?, scale_factor=?, updated_at=?
            WHERE run_id=?""",
            (r["steps_json"], r["recipe_json"], r["scaling"],
             r["scale_factor"], now, run_id))
        conn.commit()
    return {"run_id": run_id}

@router.delete("/active-runs/{run_id}")
def delete_active_run(run_id: str):
    with get_db() as conn:
        conn.execute("DELETE FROM active_runs WHERE run_id=?", (run_id,))
        conn.commit()
    return {"deleted": run_id}


# ── Daily check-in support ──────────────────────────────────────────────────
# Runs started on previous calendar days trigger a blocking popup on app load.
# The user can mark them done, snooze (don't ask again for 24h), or cancel.

@router.get("/active-runs/stale")
def list_stale_runs():
    """Return runs that should trigger the 'are you still running this?' prompt.
    Criteria: started before today (local date — but we compare ISO strings, see note)
    AND (snoozed_until IS NULL OR snoozed_until < now).

    Note on date comparison: started_at is stored as UTC ISO. We compare its
    date prefix to today's UTC date. If you're in a timezone where 'today' differs
    by hour, edge cases near midnight may misfire — acceptable trade-off for keeping
    this lightweight and avoiding a tz library."""
    now = datetime.utcnow()
    today = now.strftime("%Y-%m-%d")
    now_iso = now.isoformat()
    with get_db() as conn:
        rows = conn.execute(
            "SELECT * FROM active_runs "
            "WHERE substr(started_at, 1, 10) < ? "
            "  AND (snoozed_until IS NULL OR snoozed_until = '' OR snoozed_until < ?) "
            "ORDER BY started_at ASC",
            (today, now_iso)
        ).fetchall()
    return {"runs": [dict(r) for r in rows]}


class SnoozeRequest(BaseModel):
    hours: Optional[int] = 24


@router.post("/active-runs/{run_id}/snooze")
def snooze_active_run(run_id: str, body: SnoozeRequest):
    """Suppress the daily prompt for this run for `hours` hours.
    Used when the user clicks 'Still running' — long-running things like
    overnight cultures should be able to defer the question."""
    hours = max(1, min(body.hours or 24, 24 * 14))  # clamp to 1h..14d for sanity
    until = (datetime.utcnow() + timedelta(hours=hours)).isoformat()
    now = datetime.utcnow().isoformat()
    with get_db() as conn:
        row = conn.execute("SELECT 1 FROM active_runs WHERE run_id=?", (run_id,)).fetchone()
        if not row:
            raise HTTPException(404, "Active run not found")
        conn.execute(
            "UPDATE active_runs SET snoozed_until=?, updated_at=? WHERE run_id=?",
            (until, now, run_id)
        )
        conn.commit()
    return {"run_id": run_id, "snoozed_until": until}


class CancelRunRequest(BaseModel):
    reason: Optional[str] = ""
    date:   Optional[str] = None   # YYYY-MM-DD; defaults to today


@router.post("/active-runs/{run_id}/cancel")
def cancel_active_run(run_id: str, body: CancelRunRequest):
    """Cancel a run: save a notebook entry marked CANCELLED, delete the active
    run record. Notebook entry preserves the partial step progress + recipe so
    you can see what was done before cancellation."""
    now = datetime.utcnow().isoformat()
    date = body.date or datetime.utcnow().strftime("%Y-%m-%d")
    with get_db() as conn:
        row = conn.execute(
            "SELECT * FROM active_runs WHERE run_id=?", (run_id,)).fetchone()
        if not row:
            raise HTTPException(404, "Active run not found")
        r = dict(row)

    # Reconstruct steps + recipe for the notebook entry
    try:
        protocol = json.loads(r.get("protocol_json") or "{}")
    except Exception:
        protocol = {}
    try:
        steps = json.loads(r.get("steps_json") or "[]")
    except Exception:
        steps = []
    done_steps = [s for s in steps if s.get("done")]
    devs = [s.get("deviation", "") for s in steps if s.get("deviation")]

    # Build notes body. "[CANCELLED]" prefix makes it grep-friendly in the
    # notebook so you can find these later.
    title = protocol.get("title") or f"Protocol #{r.get('protocol_id')}"
    lines = [
        f"[CANCELLED] Protocol: {title}",
        f"Started: {(r.get('started_at') or '')[:16].replace('T', ' ')}",
        f"Cancelled: {now[:16].replace('T', ' ')}",
        f"Progress: {len(done_steps)} / {len(steps)} steps completed before cancellation",
    ]
    if body.reason:
        lines.append(f"Reason: {body.reason}")
    if done_steps:
        lines.append("\nSteps completed before cancellation:")
        for s in done_steps:
            lines.append(f"- {s.get('text', '')}")
            if s.get("deviation"):
                lines.append(f"  ↳ deviation: {s['deviation']}")
    notes_body = "\n".join(lines)

    with get_db() as conn:
        cur = conn.execute(
            "INSERT INTO entries (title,group_name,subgroup,date,notes,results,yields,issues,created,updated) "
            "VALUES (?,?,?,?,?,?,?,?,?,?)",
            (
                f"[CANCELLED] {title}",
                r.get("group_name") or "",
                r.get("subgroup") or "",
                date,
                notes_body,
                "",
                "",
                f"Run cancelled. Reason: {body.reason or '(none given)'}",
                now, now,
            ),
        )
        entry_id = cur.lastrowid
        # Also save a protocol_runs row marked cancelled — keeps the protocol
        # history honest (you can see what happened).
        conn.execute(
            "INSERT INTO protocol_runs (protocol_id,date,group_name,steps_done,steps_total,"
            "deviations,steps_json,recipe_json,entry_id,created) VALUES (?,?,?,?,?,?,?,?,?,?)",
            (
                r.get("protocol_id"),
                date,
                r.get("group_name") or "",
                len(done_steps),
                len(steps),
                "\n".join(devs) + (f"\n[CANCELLED: {body.reason or '(no reason given)'}]"),
                r.get("steps_json") or "[]",
                r.get("recipe_json"),
                entry_id,
                now,
            ),
        )
        conn.execute("DELETE FROM active_runs WHERE run_id=?", (run_id,))
        conn.commit()
    return {"cancelled": run_id, "entry_id": entry_id}

@router.post("/protocols")
async def create_from_url(body: CreateProtocol):
    now = datetime.utcnow().isoformat()
    source_text = ""
    if body.url:
        source_text = await fetch_url_text(body.url)
        if source_text and source_text.startswith("Error"):
            source_text = ""
    with get_db() as conn:
        cur = conn.execute(
            "INSERT INTO protocols (title,source_type,url,source_text,steps,recipe,notes,tags,auto_complete,created,updated) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
            (body.title, "url", body.url, source_text, None, json.dumps(DEFAULT_RECIPE), body.notes, json.dumps(body.tags), body.auto_complete, now, now))
        conn.commit()
        proto = dict(conn.execute("SELECT * FROM protocols WHERE id=?", (cur.lastrowid,)).fetchone())
    return proto

@router.post("/protocols/from-paste")
async def create_from_paste(body: PasteProtocol):
    now = datetime.utcnow().isoformat()
    with get_db() as conn:
        cur = conn.execute(
            "INSERT INTO protocols (title,source_type,url,source_text,steps,recipe,notes,tags,auto_complete,created,updated) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
            (body.title, "paste", None, body.text[:50000], None, json.dumps(DEFAULT_RECIPE), body.notes, json.dumps(body.tags), body.auto_complete, now, now))
        conn.commit()
        proto = dict(conn.execute("SELECT * FROM protocols WHERE id=?", (cur.lastrowid,)).fetchone())
    return proto

@router.post("/protocols/from-file")
async def create_from_file(
    title: str = Form(...),
    notes: str = Form(""),
    tags:  str = Form("[]"),
    auto_complete: str = Form("manual"),
    file:  UploadFile = File(...),
):
    now         = datetime.utcnow().isoformat()
    source_text = await _text_from_upload(file)
    with get_db() as conn:
        cur = conn.execute(
            "INSERT INTO protocols (title,source_type,url,source_text,steps,recipe,notes,tags,auto_complete,created,updated) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
            (title, "file", None, source_text[:50000], None, json.dumps(DEFAULT_RECIPE), notes, tags, auto_complete, now, now))
        conn.commit()
        proto = dict(conn.execute("SELECT * FROM protocols WHERE id=?", (cur.lastrowid,)).fetchone())
    return proto

@router.post("/protocols/from-manual")
async def create_manual(body: ManualProtocol):
    now    = datetime.utcnow().isoformat()
    steps  = json.dumps([{"text": s.strip()} for s in body.steps if s.strip()])
    recipe = body.recipe or json.dumps(DEFAULT_RECIPE)
    with get_db() as conn:
        cur = conn.execute(
            "INSERT INTO protocols (title,source_type,url,source_text,steps,recipe,notes,tags,auto_complete,created,updated) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
            (body.title, "manual", None, None, steps, recipe, body.notes, json.dumps(body.tags), body.auto_complete, now, now))
        conn.commit()
        return dict(conn.execute("SELECT * FROM protocols WHERE id=?", (cur.lastrowid,)).fetchone())

@router.put("/protocols/{protocol_id}")
async def update_protocol(protocol_id: int, body: UpdateProtocol):
    with get_db() as conn:
        row = conn.execute("SELECT * FROM protocols WHERE id=?", (protocol_id,)).fetchone()
        if not row:
            raise HTTPException(404, "Not found")
        p = dict(row)
        if body.title  is not None: p["title"]  = body.title
        if body.notes  is not None: p["notes"]  = body.notes
        if body.steps  is not None: p["steps"]  = body.steps
        if body.recipe is not None: p["recipe"] = body.recipe
        if body.tags   is not None: p["tags"]   = json.dumps(body.tags)
        if body.auto_complete is not None: p["auto_complete"] = body.auto_complete
        p["updated"] = datetime.utcnow().isoformat()
        conn.execute(
            "UPDATE protocols SET title=?,notes=?,steps=?,recipe=?,tags=?,auto_complete=?,updated=? WHERE id=?",
            (p["title"], p["notes"], p["steps"], p["recipe"], p["tags"], p.get("auto_complete", "manual"), p["updated"], protocol_id))
        conn.commit()
    return p

@router.delete("/protocols/{protocol_id}")
def delete_protocol(protocol_id: int):
    with get_db() as conn:
        conn.execute("DELETE FROM protocols WHERE id=?", (protocol_id,))
        conn.commit()
    return {"deleted": protocol_id}

@router.get("/protocols/{protocol_id}/claude-prompt")
def claude_prompt(protocol_id: int):
    """Return a ready-to-paste prompt for Claude chat. Frontend clipboards it."""
    with get_db() as conn:
        row = conn.execute("SELECT * FROM protocols WHERE id=?", (protocol_id,)).fetchone()
    if not row:
        raise HTTPException(404, "Not found")
    p = dict(row)
    source = _clean_source_text(p.get("source_text") or "")
    if not source:
        raise HTTPException(400, "No source text stored for this protocol")
    return {"prompt": CLAUDE_PROMPT_TEMPLATE + source}


class ImportFromClaude(BaseModel):
    formatted: str


@router.post("/protocols/{protocol_id}/import-from-claude")
def import_from_claude(protocol_id: int, body: ImportFromClaude):
    """Parse a Claude-formatted delimited block and overwrite steps + recipe;
    append the NOTES section to existing notes (separator preserves prior notes).
    Returns the updated protocol row."""
    with get_db() as conn:
        row = conn.execute("SELECT * FROM protocols WHERE id=?", (protocol_id,)).fetchone()
    if not row:
        raise HTTPException(404, "Not found")
    p = dict(row)

    sections = _split_claude_sections(body.formatted)
    if not any(sections.values()):
        raise HTTPException(400,
            "Could not find any === STEPS ===, === RECIPES ===, or === NOTES === "
            "delimiters in the pasted text")

    steps_json = _parse_steps_block(sections["STEPS"]) if sections["STEPS"] else None
    recipe_json = _parse_recipes_block(sections["RECIPES"]) if sections["RECIPES"] else None

    # Notes: append with a separator so any manual context already there survives.
    new_notes = p.get("notes") or ""
    if sections["NOTES"]:
        addition = sections["NOTES"].strip()
        if new_notes.strip():
            new_notes = new_notes.rstrip() + "\n\n---\n\n" + addition
        else:
            new_notes = addition

    now = datetime.utcnow().isoformat()
    final_steps  = steps_json  if steps_json  is not None else p.get("steps")
    final_recipe = recipe_json if recipe_json is not None else p.get("recipe")

    with get_db() as conn:
        conn.execute(
            "UPDATE protocols SET steps=?, recipe=?, notes=?, updated=? WHERE id=?",
            (final_steps, final_recipe, new_notes, now, protocol_id))
        conn.commit()
        updated = dict(conn.execute("SELECT * FROM protocols WHERE id=?", (protocol_id,)).fetchone())

    # Report what got parsed so the UI can toast a summary
    try:
        step_count = len(json.loads(final_steps or "[]"))
    except Exception:
        step_count = 0
    try:
        r = json.loads(final_recipe) if final_recipe else None
        table_count = len(r) if isinstance(r, list) else (1 if isinstance(r, dict) and r.get("columns") else 0)
    except Exception:
        table_count = 0

    return {"protocol": updated, "steps_parsed": step_count,
            "tables_parsed": table_count,
            "notes_appended": bool(sections["NOTES"].strip())}


@router.post("/protocols/{protocol_id}/clone")
def clone_protocol(protocol_id: int):
    with get_db() as conn:
        row = conn.execute("SELECT * FROM protocols WHERE id=?", (protocol_id,)).fetchone()
    if not row:
        raise HTTPException(404, "Not found")
    p = dict(row)
    now = datetime.utcnow().isoformat()
    with get_db() as conn:
        cur = conn.execute(
            "INSERT INTO protocols (title,source_type,url,source_text,steps,recipe,notes,tags,auto_complete,created,updated) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
            (p["title"] + " (copy)", p["source_type"], p["url"], p["source_text"],
             p["steps"], p["recipe"], "", p["tags"], p.get("auto_complete", "manual"), now, now))
        conn.commit()
        return dict(conn.execute("SELECT * FROM protocols WHERE id=?", (cur.lastrowid,)).fetchone())

@router.post("/active-runs/check-expiry")
def check_run_expiry():
    """Auto-complete active runs whose protocol's auto_complete window has elapsed."""
    now = datetime.utcnow()
    today = now.date().isoformat()
    completed = []
    with get_db() as conn:
        runs = conn.execute("SELECT * FROM active_runs").fetchall()
        for run_row in runs:
            run = dict(run_row)
            proto = conn.execute("SELECT auto_complete FROM protocols WHERE id=?",
                                 (run["protocol_id"],)).fetchone()
            if not proto:
                continue
            ac = (dict(proto).get("auto_complete") or "manual")
            if ac == "manual":
                continue
            started = run["started_at"][:10]
            start_date = datetime.fromisoformat(started).date() if 'T' in started else datetime.strptime(started, "%Y-%m-%d").date()
            if ac == "end_of_day":
                expiry = start_date
            else:
                # parse "Xd" format
                try:
                    days = int(ac.replace("d", ""))
                except ValueError:
                    continue
                expiry = start_date + timedelta(days=days)
            if now.date() > expiry:
                # auto-complete: save to protocol_runs, remove from active_runs
                steps = json.loads(run.get("steps_json") or "[]")
                done = sum(1 for s in steps if s.get("done"))
                devs = sum(1 for s in steps if s.get("deviation", "").strip())
                conn.execute(
                    "INSERT INTO protocol_runs (protocol_id,date,group_name,steps_done,steps_total,"
                    "deviations,steps_json,recipe_json,entry_id,created) VALUES (?,?,?,?,?,?,?,?,?,?)",
                    (run["protocol_id"], started, run.get("group_name", ""),
                     done, len(steps), devs, run.get("steps_json"),
                     run.get("recipe_json"), None, now.isoformat()))
                conn.execute("DELETE FROM active_runs WHERE run_id=?", (run["run_id"],))
                completed.append({"run_id": run["run_id"], "protocol_id": run["protocol_id"]})
        if completed:
            conn.commit()
    return {"completed": completed, "checked": len(runs) if 'runs' in dir() else 0}
