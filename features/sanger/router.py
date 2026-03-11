"""Protocols feature - protocol library with LLM extraction, manual entry, recipe tables, and run history."""
from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from pydantic import BaseModel
from typing import Optional, List
from datetime import datetime
import json, re, io

from core.database import register_table, get_db
from core.llm import llm, fetch_url_text
from core.ssh import ensure_pc_online, start_llm, call_llm_3090

register_table("protocols", """CREATE TABLE IF NOT EXISTS protocols (
    id          INTEGER PRIMARY KEY AUTOINCREMENT,
    title       TEXT NOT NULL,
    source_type TEXT NOT NULL DEFAULT 'url',
    url         TEXT DEFAULT NULL,
    source_text TEXT DEFAULT NULL,
    steps       TEXT DEFAULT NULL,
    recipe      TEXT DEFAULT NULL,
    notes       TEXT NOT NULL DEFAULT '',
    tags        TEXT NOT NULL DEFAULT '[]',
    created     TEXT NOT NULL,
    updated     TEXT NOT NULL)""")

register_table("protocol_runs", """CREATE TABLE IF NOT EXISTS protocol_runs (
    id           INTEGER PRIMARY KEY AUTOINCREMENT,
    protocol_id  INTEGER NOT NULL,
    date         TEXT NOT NULL,
    group_name   TEXT NOT NULL DEFAULT '',
    steps_done   INTEGER NOT NULL DEFAULT 0,
    steps_total  INTEGER NOT NULL DEFAULT 0,
    deviations   INTEGER NOT NULL DEFAULT 0,
    steps_json   TEXT DEFAULT NULL,
    recipe_json  TEXT DEFAULT NULL,
    entry_id     INTEGER DEFAULT NULL,
    created      TEXT NOT NULL)""")

def _migrate():
    with get_db() as conn:
        for stmt in [
            "ALTER TABLE protocols ADD COLUMN source_type TEXT NOT NULL DEFAULT 'url'",
            "ALTER TABLE protocols ADD COLUMN recipe TEXT DEFAULT NULL",
        ]:
            try:
                conn.execute(stmt)
                conn.commit()
            except Exception:
                pass
_migrate()

# --------------------------------------------------------------------------- #
#  LLM helpers
# --------------------------------------------------------------------------- #
_EXTRACT_SYSTEM = (
    "You are a lab protocol parser. Extract every step from the provided protocol text. "
    "Return ONLY a valid JSON array - no markdown fences, no explanation, nothing else. "
    "Each element must be a JSON object with a single key 'text' containing one step. "
    "Include reagents, volumes, temperatures, and timings in the step text. "
    'Example: [{"text":"Add 50 uL Buffer EB to spin column"},{"text":"Incubate 1 min at RT"}]'
)

_EXTRACT_RECIPE_SYSTEM = (
    "You are a lab protocol parser. Extract the reaction setup / reagent table from this protocol. "
    "Return a JSON object with two keys: 'columns' (array of column name strings) and "
    "'rows' (array of arrays, one per reagent). "
    "Use columns: Component, Stock conc., Volume (uL), Final conc. where applicable. "
    "Include every reagent, buffer, enzyme, and control. "
    "If no clear recipe table exists, return an object with empty rows array. "
    "Output the JSON then place /// on its own line to mark the end. "
    "No other text before the JSON."
)

def _clean_source_text(text: str) -> str:
    text = re.sub(r'<style[^>]*>.*?</style>', ' ', text, flags=re.DOTALL)
    text = re.sub(r'<script[^>]*>.*?</script>', ' ', text, flags=re.DOTALL)
    text = re.sub(r'<[^>]+>', ' ', text)
    text = re.sub(r'&nbsp;', ' ', text)
    text = re.sub(r'&[a-z]+;', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text[:4000]

def _parse_steps_raw(raw: str) -> str:
    cleaned = re.sub(r"^```[a-z]*\n?", "", raw.strip())
    cleaned = re.sub(r"\n?```$", "", cleaned.strip())
    try:
        steps = json.loads(cleaned)
        if isinstance(steps, list):
            parsed = [{"text": str(s["text"])} for s in steps if "text" in s and str(s["text"]).strip()]
            if parsed:
                return json.dumps(parsed)
    except Exception:
        pass
    steps = []
    for line in raw.splitlines():
        line = re.sub(r"^\d+[\.\)]\s*", "", line.strip())
        if line and len(line) > 4:
            steps.append({"text": line})
    return json.dumps(steps) if steps else json.dumps([])

def _parse_recipe_raw(raw: str):
    if "///" in raw:
        raw = raw[:raw.index("///")]
    cleaned = re.sub(r"^```[a-z]*\n?", "", raw.strip())
    cleaned = re.sub(r"\n?```$", "", cleaned.strip())
    m = re.search(r"\{.*\}", cleaned, re.DOTALL)
    if m:
        cleaned = m.group()
    try:
        data = json.loads(cleaned)
        if isinstance(data, dict) and "columns" in data and "rows" in data:
            cols = [str(c) for c in data["columns"] if str(c).strip()]
            rows = []
            for row in data["rows"]:
                if isinstance(row, list):
                    rows.append([str(c) for c in row])
            if cols:
                return json.dumps({"columns": cols, "rows": rows})
    except Exception:
        pass
    return None

async def extract_steps(title: str, source_text: str) -> str:
    clean_text = _clean_source_text(source_text)
    prompt = "Protocol: " + title + "\n\nText:\n" + clean_text
    try:
        if ensure_pc_online() and start_llm():
            raw = call_llm_3090(_EXTRACT_SYSTEM, prompt, max_tokens=1500)
            if raw:
                result = _parse_steps_raw(raw)
                if result != "[]":
                    return result
    except Exception:
        pass
    raw = await llm(prompt, _EXTRACT_SYSTEM, max_tokens=1500)
    return _parse_steps_raw(raw)

async def extract_recipe(title: str, source_text: str) -> str:
    clean_text = _clean_source_text(source_text)
    prompt = "Protocol: " + title + "\n\nText:\n" + clean_text
    raw = None
    try:
        if ensure_pc_online() and start_llm():
            raw = call_llm_3090(_EXTRACT_RECIPE_SYSTEM, prompt, max_tokens=800)
    except Exception:
        pass
    if not raw:
        raw = await llm(prompt, _EXTRACT_RECIPE_SYSTEM, max_tokens=800)
    result = _parse_recipe_raw(raw)
    return result if result else json.dumps(DEFAULT_RECIPE)

async def _text_from_upload(file: UploadFile) -> str:
    content = await file.read()
    fname = (file.filename or "").lower()
    if fname.endswith(".pdf"):
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            return " ".join(page.extract_text() or "" for page in reader.pages)[:6000]
        except ImportError:
            raise HTTPException(400, "pypdf not installed")
    if fname.endswith(".docx"):
        try:
            import docx
            from docx.oxml.ns import qn
            doc = docx.Document(io.BytesIO(content))
            lines = []
            for p in doc.paragraphs:
                t = p.text.strip()
                if t:
                    lines.append(t)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            for el in doc.element.body.iter():
                if el.tag == qn('w:txbxContent'):
                    for child in el.iter(qn('w:t')):
                        t = (child.text or "").strip()
                        if t:
                            lines.append(t)
            extracted = "\n".join(lines).strip()
            if not extracted:
                raise HTTPException(400, "Could not extract text from this .docx")
            return extracted[:6000]
        except HTTPException:
            raise
        except ImportError:
            raise HTTPException(400, "python-docx not installed")
        except Exception as e:
            raise HTTPException(400, "Failed to read .docx: " + str(e))
    return content.decode("utf-8", errors="ignore")[:6000]

DEFAULT_RECIPE = {
    "columns": ["Component", "Stock conc.", "Volume (uL)", "Final conc."],
    "rows": []
}

# --------------------------------------------------------------------------- #
#  Models
# --------------------------------------------------------------------------- #
class CreateProtocol(BaseModel):
    title: str
    url:   Optional[str] = None
    notes: str = ""
    tags:  List[str] = []

class PasteProtocol(BaseModel):
    title: str
    text:  str
    notes: str = ""
    tags:  List[str] = []

class ManualProtocol(BaseModel):
    title:  str
    steps:  List[str] = []
    recipe: Optional[str] = None
    notes:  str = ""
    tags:   List[str] = []

class UpdateProtocol(BaseModel):
    title:  Optional[str] = None
    notes:  Optional[str] = None
    tags:   Optional[List[str]] = None
    steps:  Optional[str] = None
    recipe: Optional[str] = None

class SaveRun(BaseModel):
    protocol_id: int
    date:        str
    group_name:  str
    steps_json:  str
    recipe_json: Optional[str] = None
    entry_id:    Optional[int] = None

# --------------------------------------------------------------------------- #
#  Router
# --------------------------------------------------------------------------- #
router = APIRouter(prefix="/api", tags=["protocols"])

@router.get("/protocols")
def list_protocols(tag: str = ""):
    with get_db() as conn:
        rows = conn.execute("SELECT * FROM protocols ORDER BY created DESC").fetchall()
    protocols = [dict(r) for r in rows]
    if tag:
        protocols = [p for p in protocols if tag in json.loads(p.get("tags") or "[]")]
    return {"protocols": protocols}

@router.get("/protocols/{protocol_id}")
def get_protocol(protocol_id: int):
    with get_db() as conn:
        row = conn.execute("SELECT * FROM protocols WHERE id=?", (protocol_id,)).fetchone()
    if not row:
        raise HTTPException(404, "Not found")
    return dict(row)

@router.get("/protocols/{protocol_id}/runs")
def get_protocol_runs(protocol_id: int):
    with get_db() as conn:
        rows = conn.execute(
            "SELECT * FROM protocol_runs WHERE protocol_id=? ORDER BY created DESC",
            (protocol_id,)).fetchall()
    return {"runs": [dict(r) for r in rows]}

@router.post("/protocol-runs")
def save_protocol_run(body: SaveRun):
    now = datetime.utcnow().isoformat()
    steps = json.loads(body.steps_json) if body.steps_json else []
    done  = sum(1 for s in steps if s.get("done"))
    devs  = sum(1 for s in steps if s.get("deviation", "").strip())
    with get_db() as conn:
        cur = conn.execute(
            "INSERT INTO protocol_runs (protocol_id,date,group_name,steps_done,steps_total,"
            "deviations,steps_json,recipe_json,entry_id,created) VALUES (?,?,?,?,?,?,?,?,?,?)",
            (body.protocol_id, body.date, body.group_name, done, len(steps),
             devs, body.steps_json, body.recipe_json, body.entry_id, now))
        conn.commit()
        return dict(conn.execute("SELECT * FROM protocol_runs WHERE id=?", (cur.lastrowid,)).fetchone())

@router.post("/protocols")
async def create_from_url(body: CreateProtocol):
    now = datetime.utcnow().isoformat()
    source_text, steps = "", None
    recipe = json.dumps(DEFAULT_RECIPE)
    if body.url:
        source_text = await fetch_url_text(body.url)
        if source_text and not source_text.startswith("Error"):
            steps  = await extract_steps(body.title, source_text)
            recipe = await extract_recipe(body.title, source_text)
    with get_db() as conn:
        cur = conn.execute(
            "INSERT INTO protocols (title,source_type,url,source_text,steps,recipe,notes,tags,created,updated) VALUES (?,?,?,?,?,?,?,?,?,?)",
            (body.title, "url", body.url, source_text, steps, recipe, body.notes, json.dumps(body.tags), now, now))
        conn.commit()
        return dict(conn.execute("SELECT * FROM protocols WHERE id=?", (cur.lastrowid,)).fetchone())

@router.post("/protocols/from-paste")
async def create_from_paste(body: PasteProtocol):
    now    = datetime.utcnow().isoformat()
    steps  = await extract_steps(body.title, body.text)
    recipe = await extract_recipe(body.title, body.text)
    with get_db() as conn:
        cur = conn.execute(
            "INSERT INTO protocols (title,source_type,url,source_text,steps,recipe,notes,tags,created,updated) VALUES (?,?,?,?,?,?,?,?,?,?)",
            (body.title, "paste", None, body.text, steps, recipe, body.notes, json.dumps(body.tags), now, now))
        conn.commit()
        return dict(conn.execute("SELECT * FROM protocols WHERE id=?", (cur.lastrowid,)).fetchone())

@router.post("/protocols/from-file")
async def create_from_file(
    title: str = Form(...),
    notes: str = Form(""),
    tags:  str = Form("[]"),
    file:  UploadFile = File(...),
):
    now         = datetime.utcnow().isoformat()
    source_text = await _text_from_upload(file)
    steps  = await extract_steps(title, source_text)
    recipe = await extract_recipe(title, source_text)
    with get_db() as conn:
        cur = conn.execute(
            "INSERT INTO protocols (title,source_type,url,source_text,steps,recipe,notes,tags,created,updated) VALUES (?,?,?,?,?,?,?,?,?,?)",
            (title, "file", None, source_text, steps, recipe, notes, tags, now, now))
        conn.commit()
        return dict(conn.execute("SELECT * FROM protocols WHERE id=?", (cur.lastrowid,)).fetchone())

@router.post("/protocols/from-manual")
async def create_manual(body: ManualProtocol):
    now    = datetime.utcnow().isoformat()
    steps  = json.dumps([{"text": s.strip()} for s in body.steps if s.strip()])
    recipe = body.recipe or json.dumps(DEFAULT_RECIPE)
    with get_db() as conn:
        cur = conn.execute(
            "INSERT INTO protocols (title,source_type,url,source_text,steps,recipe,notes,tags,created,updated) VALUES (?,?,?,?,?,?,?,?,?,?)",
            (body.title, "manual", None, None, steps, recipe, body.notes, json.dumps(body.tags), now, now))
        conn.commit()
        return dict(conn.execute("SELECT * FROM protocols WHERE id=?", (cur.lastrowid,)).fetchone())

@router.put("/protocols/{protocol_id}")
async def update_protocol(protocol_id: int, body: UpdateProtocol):
    with get_db() as conn:
        row = conn.execute("SELECT * FROM protocols WHERE id=?", (protocol_id,)).fetchone()
        if not row:
            raise HTTPException(404, "Not found")
        p = dict(row)
        if body.title  is not None: p["title"]  = body.title
        if body.notes  is not None: p["notes"]  = body.notes
        if body.steps  is not None: p["steps"]  = body.steps
        if body.recipe is not None: p["recipe"] = body.recipe
        if body.tags   is not None: p["tags"]   = json.dumps(body.tags)
        p["updated"] = datetime.utcnow().isoformat()
        conn.execute(
            "UPDATE protocols SET title=?,notes=?,steps=?,recipe=?,tags=?,updated=? WHERE id=?",
            (p["title"], p["notes"], p["steps"], p["recipe"], p["tags"], p["updated"], protocol_id))
        conn.commit()
    return p

@router.delete("/protocols/{protocol_id}")
def delete_protocol(protocol_id: int):
    with get_db() as conn:
        conn.execute("DELETE FROM protocols WHERE id=?", (protocol_id,))
        conn.commit()
    return {"deleted": protocol_id}

@router.post("/protocols/{protocol_id}/re-extract")
async def re_extract_steps(protocol_id: int):
    with get_db() as conn:
        row = conn.execute("SELECT * FROM protocols WHERE id=?", (protocol_id,)).fetchone()
    if not row:
        raise HTTPException(404, "Not found")
    p = dict(row)
    if not p["source_text"]:
        raise HTTPException(400, "No source text stored - re-import the protocol")
    steps = await extract_steps(p["title"], p["source_text"])
    with get_db() as conn:
        conn.execute("UPDATE protocols SET steps=?, updated=? WHERE id=?",
                     (steps, datetime.utcnow().isoformat(), protocol_id))
        conn.commit()
    return {"steps": steps}
